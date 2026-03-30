"""
RAG Engine for AI OS
Handles document loading, embedding, and retrieval
"""

import hashlib
import logging
import os
import pickle
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
import time

import faiss
import numpy as np
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer

from backend.models import SearchResult, SearchRequest, IndexResponse

logger = logging.getLogger(__name__)


class Document:
    """Represents a document chunk"""
    
    def __init__(
        self,
        content: str,
        source: str,
        chunk_id: int = 0,
        page: Optional[int] = None,
        metadata: Optional[Dict] = None,
    ):
        self.content = content
        self.source = source
        self.chunk_id = chunk_id
        self.page = page
        self.metadata = metadata or {}
        self.embedding: Optional[np.ndarray] = None
    
    def __repr__(self) -> str:
        return f"Document(source={self.source}, chunk={self.chunk_id}, len={len(self.content)})"


class RAGEngine:
    """Retrieval Augmented Generation Engine"""
    
    def __init__(
        self,
        embedding_model: str = "all-MiniLM-L6-v2",
        device: str = "cpu",
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        vector_store_path: str = "./data/vector_store",
        top_k: int = 5,
        similarity_threshold: float = 0.7,
    ):
        """
        Initialize RAG engine
        
        Args:
            embedding_model: Sentence transformer model name
            device: Device for embeddings (cpu/cuda)
            chunk_size: Document chunk size
            chunk_overlap: Chunk overlap
            vector_store_path: Path to store vector index
            top_k: Number of results to retrieve
            similarity_threshold: Minimum similarity score
        """
        self.embedding_model_name = embedding_model
        self.device = device
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.vector_store_path = Path(vector_store_path)
        self.top_k = top_k
        self.similarity_threshold = similarity_threshold
        
        # Create vector store directory
        self.vector_store_path.mkdir(parents=True, exist_ok=True)
        
        # Initialize embedding model (lazy loading)
        self._embedding_model: Optional[SentenceTransformer] = None
        
        # Initialize text splitter
        self._text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        
        # Document store
        self._documents: Dict[int, Document] = {}
        self._next_id = 0
        
        # FAISS index
        self._index: Optional[faiss.Index] = None
        self._embedding_dim: Optional[int] = None
        
        logger.info(f"RAG engine initialized with model: {embedding_model}")
    
    @property
    def embedding_model(self) -> SentenceTransformer:
        """Lazy load embedding model"""
        if self._embedding_model is None:
            logger.info(f"Loading embedding model: {self.embedding_model_name}")
            self._embedding_model = SentenceTransformer(
                self.embedding_model_name,
                device=self.device,
            )
        return self._embedding_model
    
    def _get_embedding(self, text: str) -> np.ndarray:
        """
        Get embedding for text
        
        Args:
            text: Text to embed
            
        Returns:
            Embedding vector
        """
        embedding = self.embedding_model.encode(text, convert_to_numpy=True)
        return embedding.astype(np.float32)
    
    def _chunk_document(self, content: str, source: str) -> List[Document]:
        """
        Split document into chunks
        
        Args:
            content: Document content
            source: Document source path
            
        Returns:
            List of document chunks
        """
        chunks = self._text_splitter.split_text(content)
        
        documents = []
        for i, chunk in enumerate(chunks):
            doc = Document(
                content=chunk,
                source=source,
                chunk_id=i,
            )
            documents.append(doc)
        
        return documents
    
    def _read_file(self, file_path: str) -> Optional[str]:
        """
        Read file content
        
        Args:
            file_path: Path to file
            
        Returns:
            File content or None if failed
        """
        path = Path(file_path)
        
        if not path.exists():
            logger.error(f"File not found: {file_path}")
            return None
        
        try:
            # Handle different file types
            suffix = path.suffix.lower()
            
            if suffix == ".pdf":
                return self._read_pdf(file_path)
            elif suffix in [".docx", ".doc"]:
                return self._read_docx(file_path)
            else:
                # Text files
                with open(path, "r", encoding="utf-8", errors="replace") as f:
                    return f.read()
                    
        except Exception as e:
            logger.error(f"Error reading file {file_path}: {e}")
            return None
    
    def _read_pdf(self, file_path: str) -> Optional[str]:
        """Read PDF file"""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except ImportError:
            logger.error("pypdf not installed, cannot read PDF")
            return None
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
            return None
    
    def _read_docx(self, file_path: str) -> Optional[str]:
        """Read DOCX file"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except ImportError:
            logger.error("python-docx not installed, cannot read DOCX")
            return None
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            return None
    
    def index_file(self, file_path: str) -> int:
        """
        Index a single file
        
        Args:
            file_path: Path to file
            
        Returns:
            Number of chunks indexed
        """
        content = self._read_file(file_path)
        if content is None:
            return 0
        
        # Chunk document
        chunks = self._chunk_document(content, file_path)
        
        # Add chunks to index
        for doc in chunks:
            doc_id = self._next_id
            self._documents[doc_id] = doc
            self._next_id += 1
        
        logger.info(f"Indexed {file_path}: {len(chunks)} chunks")
        return len(chunks)
    
    def index_directory(
        self,
        directory: str,
        recursive: bool = True,
        file_types: Optional[List[str]] = None,
    ) -> IndexResponse:
        """
        Index all files in a directory
        
        Args:
            directory: Directory path
            recursive: Index subdirectories
            file_types: File extensions to include
            
        Returns:
            IndexResponse with results
        """
        start_time = time.time()
        
        dir_path = Path(directory)
        if not dir_path.exists():
            return IndexResponse(
                success=False,
                errors=[f"Directory not found: {directory}"],
            )
        
        # Default file types
        if file_types is None:
            file_types = [
                ".txt", ".md", ".markdown", ".py", ".js", ".ts", ".jsx", ".tsx",
                ".java", ".c", ".cpp", ".h", ".hpp", ".go", ".rs", ".rb", ".php",
                ".sh", ".bash", ".zsh", ".json", ".yaml", ".yml", ".xml",
                ".html", ".htm", ".css", ".scss", ".sass", ".sql", ".pdf", ".docx",
            ]
        
        # Normalize extensions
        file_types = [ext.lower() if ext.startswith(".") else f".{ext}" for ext in file_types]
        
        # Find files
        files = []
        if recursive:
            for ext in file_types:
                files.extend(dir_path.rglob(f"*{ext}"))
        else:
            for ext in file_types:
                files.extend(dir_path.glob(f"*{ext}"))
        
        # Index files
        total_chunks = 0
        errors = []
        
        for file_path in files:
            try:
                chunks = self.index_file(str(file_path))
                total_chunks += chunks
            except Exception as e:
                error_msg = f"Error indexing {file_path}: {str(e)}"
                logger.error(error_msg)
                errors.append(error_msg)
        
        # Build index
        self._build_index()
        
        duration = time.time() - start_time
        
        logger.info(f"Indexed {len(files)} files, {total_chunks} chunks in {duration:.2f}s")
        
        return IndexResponse(
            success=True,
            files_indexed=len(files),
            chunks_created=total_chunks,
            errors=errors,
            duration_seconds=duration,
        )
    
    def _build_index(self) -> None:
        """Build FAISS index from documents"""
        if not self._documents:
            logger.warning("No documents to index")
            return
        
        # Get embedding dimension from first document
        sample_doc = list(self._documents.values())[0]
        sample_embedding = self._get_embedding(sample_doc.content)
        self._embedding_dim = len(sample_embedding)
        
        # Create FAISS index
        self._index = faiss.IndexFlatIP(self._embedding_dim)  # Inner product for cosine similarity
        
        # Add all documents
        embeddings = []
        for doc_id, doc in self._documents.items():
            embedding = self._get_embedding(doc.content)
            doc.embedding = embedding
            embeddings.append(embedding)
        
        if embeddings:
            embeddings_array = np.array(embeddings)
            # Normalize for cosine similarity
            faiss.normalize_L2(embeddings_array)
            self._index.add(embeddings_array)
        
        logger.info(f"Built FAISS index with {len(embeddings)} vectors")
    
    def search(self, request: SearchRequest) -> List[SearchResult]:
        """
        Search for relevant documents
        
        Args:
            request: Search request
            
        Returns:
            List of search results
        """
        if self._index is None or not self._documents:
            logger.warning("No index available for search")
            return []
        
        # Get query embedding
        query_embedding = self._get_embedding(request.query)
        query_embedding = query_embedding.reshape(1, -1)
        faiss.normalize_L2(query_embedding)
        
        # Search
        scores, indices = self._index.search(query_embedding, request.top_k)
        
        # Build results
        results = []
        for score, idx in zip(scores[0], indices[0]):
            if idx == -1:  # FAISS returns -1 for empty slots
                continue
            
            # Convert inner product to cosine similarity (both normalized)
            similarity = float(score)
            
            if similarity < request.threshold:
                continue
            
            doc = self._documents.get(int(idx))
            if doc:
                results.append(SearchResult(
                    content=doc.content,
                    source=doc.source,
                    score=similarity,
                    page=doc.page,
                ))
        
        return results
    
    def search_simple(self, query: str, top_k: int = 5) -> List[SearchResult]:
        """
        Simple search interface
        
        Args:
            query: Search query
            top_k: Number of results
            
        Returns:
            List of search results
        """
        request = SearchRequest(query=query, top_k=top_k)
        return self.search(request)
    
    def get_context_string(self, query: str, top_k: int = 5) -> str:
        """
        Get context string for RAG
        
        Args:
            query: Search query
            top_k: Number of results
            
        Returns:
            Concatenated context string
        """
        results = self.search_simple(query, top_k)
        
        if not results:
            return "No relevant documents found."
        
        context_parts = []
        for i, result in enumerate(results, 1):
            context_parts.append(
                f"[Document {i}] Source: {result.source}\n{result.content}\n"
            )
        
        return "\n".join(context_parts)
    
    def save_index(self, name: str = "default") -> str:
        """
        Save index to disk
        
        Args:
            name: Index name
            
        Returns:
            Path to saved index
        """
        index_path = self.vector_store_path / f"{name}.faiss"
        docs_path = self.vector_store_path / f"{name}_docs.pkl"
        
        # Save FAISS index
        if self._index:
            faiss.write_index(self._index, str(index_path))
        
        # Save documents
        with open(docs_path, "wb") as f:
            pickle.dump(self._documents, f)
        
        logger.info(f"Saved index to {index_path}")
        return str(index_path)
    
    def load_index(self, name: str = "default") -> bool:
        """
        Load index from disk
        
        Args:
            name: Index name
            
        Returns:
            True if successful
        """
        index_path = self.vector_store_path / f"{name}.faiss"
        docs_path = self.vector_store_path / f"{name}_docs.pkl"
        
        if not index_path.exists() or not docs_path.exists():
            logger.warning(f"Index not found: {name}")
            return False
        
        try:
            # Load FAISS index
            self._index = faiss.read_index(str(index_path))
            self._embedding_dim = self._index.d
            
            # Load documents
            with open(docs_path, "rb") as f:
                self._documents = pickle.load(f)
            
            # Update next_id
            if self._documents:
                self._next_id = max(self._documents.keys()) + 1
            
            logger.info(f"Loaded index from {index_path} ({len(self._documents)} documents)")
            return True
            
        except Exception as e:
            logger.error(f"Error loading index: {e}")
            return False
    
    def clear_index(self) -> None:
        """Clear all indexed documents"""
        self._documents.clear()
        self._next_id = 0
        self._index = None
        logger.info("Index cleared")
    
    def get_stats(self) -> Dict[str, Any]:
        """
        Get index statistics
        
        Returns:
            Statistics dictionary
        """
        return {
            "total_documents": len(self._documents),
            "index_built": self._index is not None,
            "embedding_model": self.embedding_model_name,
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
        }


# Global RAG engine instance
_rag_engine: Optional[RAGEngine] = None


def get_rag_engine(
    embedding_model: str = "all-MiniLM-L6-v2",
    **kwargs,
) -> RAGEngine:
    """
    Get or create global RAG engine
    
    Args:
        embedding_model: Embedding model name
        **kwargs: Additional options
        
    Returns:
        RAGEngine instance
    """
    global _rag_engine
    if _rag_engine is None:
        _rag_engine = RAGEngine(embedding_model=embedding_model, **kwargs)
    return _rag_engine
